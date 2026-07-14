from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from pathlib import Path
OUT=Path('output'); OUT.mkdir(exist_ok=True)
d=Document(); s=d.sections[0]; s.top_margin=Inches(.65); s.bottom_margin=Inches(.65); s.left_margin=Inches(.75); s.right_margin=Inches(.75)
for n,z,c in [('Normal',10.5,'334155'),('Title',25,'17324D'),('Heading 1',17,'0E7490'),('Heading 2',12.5,'17324D')]:
 st=d.styles[n]; st.font.name='Malgun Gothic'; st._element.rPr.rFonts.set(qn('w:eastAsia'),'Malgun Gothic'); st.font.size=Pt(z); st.font.color.rgb=RGBColor.from_string(c); st.font.bold=n!='Normal'
def shade(c,f):
 x=c._tc.get_or_add_tcPr(); q=OxmlElement('w:shd'); q.set(qn('w:fill'),f); x.append(q)
def bullet(xs):
 for x in xs: d.add_paragraph(x,style='List Bullet')
def steps(xs):
 for x in xs: d.add_paragraph(x,style='List Number')
def note(a,b):
 t=d.add_table(rows=1,cols=1); c=t.cell(0,0); shade(c,'E8F4F7'); c.text=a+'\n'+b
def table(headers, rows):
 t=d.add_table(rows=1,cols=len(headers)); t.style='Table Grid'; t.alignment=WD_TABLE_ALIGNMENT.CENTER
 for i,h in enumerate(headers): t.cell(0,i).text=h; shade(t.cell(0,i),'0E7490')
 for row in rows:
  cs=t.add_row().cells
  for i,v in enumerate(row): cs[i].text=v
p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.add_run('JT Academy').bold=True
p=d.add_paragraph('소방설비전기기사 학습 앱\n사용설명서',style='Title'); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
p=d.add_paragraph('배포용 안내서 | MVP 버전'); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
note('이 문서의 목적','처음 사용하는 학습자가 앱을 열고 오늘 학습을 시작하여 진도와 복습 기록을 관리할 수 있도록 안내합니다.')
d.add_heading('1. 앱 소개',1); d.add_paragraph('JT Academy 소방설비전기기사 학습 앱은 개념 학습, 빈칸·암기 확인, 기출문제 풀이, 오답 복습, 모의시험, 학습현황을 하나의 흐름으로 제공하는 웹 앱입니다.')
bullet(['학습 순서: 개념 학습 → 암기 확인 → 해당 lesson 기출 풀이','오늘 학습과 오늘 복습을 분리하여 관리','오답과 기출 정답 기록을 바탕으로 복습 시점 자동 관리','학습 기록은 브라우저에 저장되며, 로그인 시 Firebase 계정과 동기화 가능'])
d.add_heading('2. 시작하기',1); d.add_heading('2.1 배포된 사이트 열기',2); d.add_paragraph('배포 담당자가 제공한 웹 주소를 최신 브라우저에서 엽니다. JavaScript가 켜져 있어야 하며, 처음 접속할 때 학습 JSON 파일을 자동으로 불러옵니다.')
d.add_heading('2.2 로컬에서 실행하기 (운영자용)',2); steps(['프로젝트 폴더에서 터미널을 엽니다.','python -m http.server 8080 명령을 실행합니다.','브라우저에서 http://localhost:8080 을 엽니다.']); note('중요','HTML 파일을 직접 더블클릭(file://)하면 JSON fetch가 차단될 수 있습니다. 반드시 웹 서버를 통해 실행하세요.')
d.add_heading('3. 권장 학습 흐름',1); d.add_paragraph('홈 화면의 오늘의 학습 시작 버튼을 누르면 저장된 진행 상태를 기준으로 다음 미완료 단계부터 이어서 진행합니다.')
steps(['오늘의 학습에서 현재 과목과 lesson을 확인합니다.','개념 학습을 읽고 개념 학습 완료를 누릅니다.','암기 문제를 모두 풀고 정답을 확인합니다.','해당 lesson의 기출문제를 풀고 채점합니다.','다음 lesson으로 이동하거나 과목을 완료합니다.','오늘 복습할 기출이 있다면 복습 학습을 이어서 진행합니다.']); note('암기 단계 안내','암기 오답은 기출 복습 큐에 추가되지 않습니다. 암기 문항을 모두 확인한 뒤 해당 lesson 기출 풀이로 진행합니다.')
d.add_heading('4. 메뉴별 사용법',1)
table(['메뉴','할 수 있는 일'],[('홈 / 학습 대시보드','오늘 신규 학습량, 오늘 복습량, 현재 과목, 진행률을 확인하고 학습을 시작합니다.'),('개념학습','lesson의 핵심 개념을 확인한 뒤 완료를 눌러 다음 단계로 이동합니다.'),('암기학습','선택형, OX, 주관식 암기 문항을 풉니다.'),('기출문제','답을 제출하고 채점 결과와 해설을 확인합니다.'),('오답노트','틀린 기출과 누적 오답을 확인하고 바로 복습합니다.'),('모의시험','시험을 시작하고 제출하여 결과와 오답을 확인합니다.'),('학습현황','과목별 진도, due 복습 수량, 취약 영역을 확인합니다.'),('설정','30일/60일 플랜과 학습 기록 초기화를 관리합니다.')])
d.add_heading('5. 복습과 오답 관리',1); d.add_paragraph('기출문제의 정답·오답 결과는 복습 큐에 반영됩니다. 홈 화면의 오늘 복습 학습 또는 오답노트에서 due 문항을 확인할 수 있습니다.')
bullet(['정답 문항은 다음 복습일이 뒤로 조정됩니다.','오답 문항은 우선순위가 올라가고 가까운 복습일로 다시 예약됩니다.','오늘 복습 세션은 시작 시점의 문항 수를 기준으로 진행률을 표시합니다.','복습 목록의 바로 풀기를 누르면 선택한 문항부터 이어서 풉니다.'])
d.add_heading('6. 학습 기록과 로그인',1); d.add_paragraph('비로그인 상태에서는 화면을 볼 수 있지만 기록을 변경하는 동작에 로그인 안내가 표시될 수 있습니다. 로그인하면 사용자별 학습 진행률을 Firebase에 동기화할 수 있습니다.')
bullet(['로그인 전: 브라우저 localStorage 기반의 로컬 학습 기록','로그인 후: 사용자 진행률을 Firestore에 저장하고 재로그인 시 복원','로그아웃: 계정 동기화 상태를 종료합니다.']); note('기록 초기화 주의','설정에서 기록을 초기화하면 진도, 오답, 복습 큐, 북마크, 모의시험 기록이 함께 삭제됩니다.')
d.add_heading('7. 배포 전 점검표',1); bullet(['배포 주소에서 홈 화면이 정상적으로 열리는가?','개념 → 암기 → 기출 순서가 정상적으로 이어지는가?','새로고침 후 진도와 오늘 학습 세션이 유지되는가?','오답이 오답노트와 복습 큐에 표시되는가?','모의시험 시작·제출·결과 확인이 가능한가?','30일/60일 플랜과 기록 초기화가 정상 동작하는가?','모바일 화면에서 버튼과 문제 내용이 잘 보이는가?','Firebase 사용 시 로그인, 로그아웃, 진행률 복원이 가능한가?'])
d.add_heading('8. 문제 해결',1)
table(['증상','확인할 사항'],[('화면이 비어 있음','웹 서버로 접속했는지 확인하고 새로고침합니다. Console과 Network에서 JSON 오류를 확인합니다.'),('문제가 로드되지 않음','data/fire_subjects.json, fire_lessons.json, fire_questions.json, fire_mock-exam.json이 배포물에 포함되었는지 확인합니다.'),('진도가 사라짐','같은 브라우저·기기인지 확인하고, 로그인 사용자는 동일 계정으로 로그인합니다.'),('수식이 이상함','인터넷 연결과 MathJax CDN 로딩 상태를 확인한 뒤 새로고침합니다.'),('로그인이 안 됨','Firebase Email/Password 제공자와 배포 도메인 허용 설정을 확인합니다.')])
d.add_paragraph('문의 시에는 접속 주소, 사용 브라우저, 발생 화면, 오류 메시지, 발생 시각을 함께 전달하면 원인 확인이 빠릅니다.')
for sec in d.sections:
 f=sec.footer.paragraphs[0]; f.alignment=WD_ALIGN_PARAGRAPH.CENTER; f.add_run('JT Academy | 소방설비전기기사 학습 앱 사용설명서').font.size=Pt(8)
d.save(OUT/'JT_Academy_소방설비전기기사_사용설명서.docx')

